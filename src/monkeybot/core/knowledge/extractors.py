"""Extract indexable text from workspace / note files (text + media)."""

from __future__ import annotations

import hashlib
import logging
import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from monkeybot.core.lockfile_names import LOCKFILE_NAMES

# Keep in sync with workspace grep ignores; also skip .monkeybot under workspace.
IGNORE_DIRS = frozenset(
    {
        ".git",
        ".hg",
        ".svn",
        "node_modules",
        "__pycache__",
        ".venv",
        "venv",
        ".mypy_cache",
        ".pytest_cache",
        ".ruff_cache",
        ".tox",
        "dist",
        "build",
        ".next",
        ".monkeybot",
    }
)

_TEXT_SUFFIXES = frozenset(
    {
        ".md",
        ".markdown",
        ".txt",
        ".rst",
        ".py",
        ".ts",
        ".tsx",
        ".js",
        ".jsx",
        ".mjs",
        ".cjs",
        ".json",
        ".yaml",
        ".yml",
        ".toml",
        ".ini",
        ".cfg",
        ".css",
        ".scss",
        ".html",
        ".htm",
        ".xml",
        ".svg",
        ".sh",
        ".bash",
        ".zsh",
        ".rs",
        ".go",
        ".java",
        ".kt",
        ".swift",
        ".c",
        ".h",
        ".cpp",
        ".hpp",
        ".cs",
        ".rb",
        ".php",
        ".sql",
        ".graphql",
        ".vue",
        ".svelte",
    }
)

PDF_SUFFIXES = frozenset({".pdf"})
DOCX_SUFFIXES = frozenset({".docx"})
IMAGE_SUFFIXES = frozenset({".png", ".jpg", ".jpeg", ".gif", ".webp"})
MEDIA_SUFFIXES = PDF_SUFFIXES | DOCX_SUFFIXES | IMAGE_SUFFIXES

logger = logging.getLogger(__name__)

_HTML_TAG_RE = re.compile(r"<[^>]+>")
_NULL_BYTE = b"\x00"

_SKIP_FILE_NAMES = LOCKFILE_NAMES


@dataclass(frozen=True)
class PdfPage:
    """One PDF page with extracted text (1-based page number)."""

    page: int
    text: str


@dataclass(frozen=True)
class ExtractedDocument:
    """Structured extraction result for a non-plain-text file."""

    kind: str  # "pdf" | "docx"
    text: str
    pages: tuple[PdfPage, ...] = ()


def content_hash(data: bytes | str) -> str:
    raw = data.encode("utf-8") if isinstance(data, str) else data
    return hashlib.sha256(raw).hexdigest()


def media_content_digest(raw: bytes, *, chunker_version: int) -> str:
    """Digest for binary media files (hash of bytes + chunker version)."""
    return content_hash(f"chunker:{chunker_version}\n".encode() + raw)


def _file_size_ok(path: Path, *, max_file_bytes: int) -> bool:
    try:
        if not path.is_file():
            return False
        size = path.stat().st_size
    except OSError:
        return False
    return 0 < size <= max_file_bytes


def is_probably_text(path: Path, *, max_file_bytes: int) -> bool:
    """True when the path looks like a text file under size limits."""
    if not _file_size_ok(path, max_file_bytes=max_file_bytes):
        return False
    suffix = path.suffix.lower()
    name = path.name.lower()
    if name == ".env.example" or suffix in _TEXT_SUFFIXES:
        return True
    if suffix in MEDIA_SUFFIXES:
        return False
    try:
        sample = path.read_bytes()[:8192]
    except OSError:
        return False
    return _NULL_BYTE not in sample


def is_media_file(path: Path, *, max_file_bytes: int) -> bool:
    """True when path is a known PDF/DOCX/image under the size limit."""
    if not _file_size_ok(path, max_file_bytes=max_file_bytes):
        return False
    return path.suffix.lower() in MEDIA_SUFFIXES


def is_indexable_file(path: Path, *, max_file_bytes: int) -> bool:
    """True for text or supported media files under size limits."""
    return is_probably_text(path, max_file_bytes=max_file_bytes) or is_media_file(
        path, max_file_bytes=max_file_bytes
    )


def read_text_file(path: Path, *, max_file_bytes: int) -> str | None:
    """Read UTF-8 text (with replacement) or return None if unreadable / binary."""
    if not is_probably_text(path, max_file_bytes=max_file_bytes):
        return None
    try:
        raw = path.read_bytes()
    except OSError:
        return None
    if len(raw) > max_file_bytes or _NULL_BYTE in raw:
        return None
    text = raw.decode("utf-8", errors="replace")
    if path.suffix.lower() in {".html", ".htm"}:
        text = strip_html(text)
    return text


def strip_html(html: str) -> str:
    """Lightweight HTML → text (tags stripped; entities left as-is)."""
    no_script = re.sub(
        r"<(script|style)[^>]*>.*?</\1>",
        " ",
        html,
        flags=re.IGNORECASE | re.DOTALL,
    )
    return _HTML_TAG_RE.sub(" ", no_script)


def extract_pdf_pages(path: Path, *, max_file_bytes: int) -> list[PdfPage] | None:
    """Extract per-page text from a PDF. Soft-fails when pypdf is missing/corrupt."""
    if not _file_size_ok(path, max_file_bytes=max_file_bytes):
        return None
    if path.suffix.lower() not in PDF_SUFFIXES:
        return None
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning(
            "knowledge PDF extract skipped for %s — install with: uv sync --extra knowledge-media",
            path,
        )
        return None
    pages: list[PdfPage] = []
    try:
        # Own the handle so it is always closed — pypdf keeps the stream open
        # for lazy page access, which leaks descriptors over a large scan.
        with path.open("rb") as stream:
            reader = PdfReader(stream)
            for i, page in enumerate(reader.pages, start=1):
                try:
                    text = (page.extract_text() or "").strip()
                except Exception as exc:
                    logger.warning(
                        "knowledge PDF page %d extract failed for %s: %r", i, path, exc
                    )
                    continue
                if text:
                    pages.append(PdfPage(page=i, text=text))
    except Exception as exc:
        logger.warning("knowledge PDF open failed for %s: %r", path, exc)
        return None
    return pages


def extract_docx_text(path: Path, *, max_file_bytes: int) -> str | None:
    """Extract paragraph *and* table text from a DOCX, in document order.

    Soft-fails when python-docx is missing.
    """
    if not _file_size_ok(path, max_file_bytes=max_file_bytes):
        return None
    if path.suffix.lower() not in DOCX_SUFFIXES:
        return None
    try:
        from docx import Document
    except ImportError:
        logger.warning(
            "knowledge DOCX extract skipped for %s — install with: uv sync --extra knowledge-media",
            path,
        )
        return None
    try:
        doc = Document(str(path))
    except Exception as exc:
        logger.warning("knowledge DOCX open failed for %s: %r", path, exc)
        return None
    try:
        parts = _docx_body_texts(doc)
    except Exception as exc:
        logger.warning(
            "knowledge DOCX body walk failed for %s: %r; falling back to paragraphs",
            path,
            exc,
        )
        parts = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    text = "\n\n".join(parts).strip()
    return text or None


def _docx_body_texts(doc: Any) -> list[str]:
    """Body paragraphs and tables in document order."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    out: list[str] = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = Paragraph(child, doc).text.strip()
            if text:
                out.append(text)
        elif child.tag == qn("w:tbl"):
            rendered = _docx_table_text(Table(child, doc))
            if rendered:
                out.append(rendered)
    return out


def _docx_table_text(table: Any) -> str:
    """Render a table as newline-separated ``cell | cell`` rows."""
    rows: list[str] = []
    for row in table.rows:
        cells = [_docx_cell_text(cell) for cell in row.cells]
        if any(cells):
            rows.append(" | ".join(cells))
    return "\n".join(rows)


def _docx_cell_text(cell: Any) -> str:
    parts = [p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()]
    for nested in cell.tables:
        rendered = _docx_table_text(nested)
        if rendered:
            parts.append(rendered)
    return " ".join(parts).strip()


def extract_file(path: Path, *, max_file_bytes: int) -> ExtractedDocument | None:
    """Extract structured text from PDF or DOCX. Images are handled by captions."""
    suffix = path.suffix.lower()
    if suffix in PDF_SUFFIXES:
        pages = extract_pdf_pages(path, max_file_bytes=max_file_bytes)
        if not pages:
            return None
        joined = "\n\n".join(f"[PDF page {p.page}]\n{p.text}" for p in pages)
        return ExtractedDocument(kind="pdf", text=joined, pages=tuple(pages))
    if suffix in DOCX_SUFFIXES:
        text = extract_docx_text(path, max_file_bytes=max_file_bytes)
        if not text:
            return None
        return ExtractedDocument(kind="docx", text=text)
    return None


def walk_indexable_files(root: Path, *, max_file_bytes: int) -> list[Path]:
    """Walk ``root`` yielding text + supported media files, skipping ignore dirs."""
    if not root.is_dir():
        return []
    out: list[Path] = []
    for dirpath, dirnames, filenames in os.walk(root):
        dirnames[:] = sorted(d for d in dirnames if d not in IGNORE_DIRS and not d.startswith("."))
        for name in sorted(filenames):
            if name.startswith(".") and name != ".env.example":
                continue
            if name.lower() in _SKIP_FILE_NAMES:
                continue
            path = Path(dirpath) / name
            if is_indexable_file(path, max_file_bytes=max_file_bytes):
                out.append(path)
    return out


def walk_text_files(root: Path, *, max_file_bytes: int) -> list[Path]:
    """Walk ``root`` yielding text files only (alias kept for callers/tests)."""
    return [
        p
        for p in walk_indexable_files(root, max_file_bytes=max_file_bytes)
        if is_probably_text(p, max_file_bytes=max_file_bytes)
    ]


__all__ = [
    "DOCX_SUFFIXES",
    "ExtractedDocument",
    "IGNORE_DIRS",
    "IMAGE_SUFFIXES",
    "MEDIA_SUFFIXES",
    "PDF_SUFFIXES",
    "PdfPage",
    "content_hash",
    "extract_docx_text",
    "extract_file",
    "extract_pdf_pages",
    "is_indexable_file",
    "is_media_file",
    "is_probably_text",
    "media_content_digest",
    "read_text_file",
    "strip_html",
    "walk_indexable_files",
    "walk_text_files",
]
